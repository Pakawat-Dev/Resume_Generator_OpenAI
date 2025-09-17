"""
Generate a one-page resume (DOCX) from a job title using OpenAI (GPT-4.1).
- Token-efficient prompt & low temperature
- Exactly 3 bullets per required section
- Two-column DOCX layout for consistency
"""

import json
import os
import sys
from datetime import datetime
from typing import Dict, List

from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Configuration constants (easier to modify in one place)
API_KEY = os.getenv("OPENAI_API_KEY")
MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
NAME = os.getenv("CANDIDATE_NAME", "Your Name")
LEVEL = os.getenv("TARGET_SENIORITY", "Senior")
CONTEXT = os.getenv("INDUSTRY_CONTEXT", "AI engineering for medical software")

# Exit early if no API key (fail fast principle)
if not API_KEY:
    sys.exit("Missing OPENAI_API_KEY in .env")

# Initialize API client once (reuse connection)
client = OpenAI(api_key=API_KEY)

# Section configuration (single source of truth)
SECTIONS = [
    "career_overview",
    "summary_profile", 
    "core_technical_competencies",
    "work_experience",
    "academic_history",
    "contact_information",
    "professional_development",
    "technical_skills_matrix"
]


def build_prompt(job_title: str, industry: str, seniority: str) -> str:
    """Build minimal prompt to save tokens."""
    # JSON schema built programmatically (easier to maintain)
    schema = {
        "type": "object",
        "properties": {
            "name": {"type": "string"},
            **{  # Dictionary comprehension (efficient)
                s: {
                    "type": "array",
                    "items": {"type": "string"},
                    "minItems": 3,
                    "maxItems": 3
                } 
                for s in SECTIONS
            }
        },
        "required": ["name"] + SECTIONS,
        "additionalProperties": False
    }
    
    # Compact prompt format (saves tokens)
    return (
        f'Generate resume JSON for "{job_title}" level "{seniority}" '
        f'industry: {industry}. Name: "{NAME}". '
        f'Schema: {json.dumps(schema, separators=(",", ":"))} '  # No spaces in JSON
        'Rules: 3 bullets each, <=16 words, achievement-focused, '
        'no markdown, ASCII quotes only.'
    )


def call_openai(job_title: str, industry: str, seniority: str) -> Dict:
    """Call API with automatic retry on parse failure."""
    prompt = build_prompt(job_title, industry, seniority)
    
    # Try twice with progressively stricter prompts
    for attempt in range(2):
        if attempt == 1:
            prompt += " OUTPUT ONLY RAW JSON."  # Stricter on retry
            
        try:
            response = client.chat.completions.create(
                model=MODEL,
                max_tokens=700,  # Limit cost
                temperature=0.1,  # More deterministic
                messages=[{"role": "user", "content": prompt}]
            )
        except Exception as e:
            if "connection" in str(e).lower() or "network" in str(e).lower():
                raise RuntimeError(f"Connection error: Check your internet connection and API key. Details: {e}")
            else:
                raise RuntimeError(f"API error: {e}")
        
        # Extract text content efficiently
        raw = response.choices[0].message.content.strip()
        
        # Clean markdown fences if present
        if raw.startswith("```"):
            raw = raw.strip("`\n")
            if raw.lower().startswith("json"):
                raw = raw[4:]
            raw = raw.rstrip("`").strip()
        
        # Try parsing
        try:
            data = json.loads(raw)
            validate_payload(data)  # Validate structure
            return data
        except (json.JSONDecodeError, ValueError) as e:
            if attempt == 1:  # Final attempt failed
                raise RuntimeError(f"Parse failed: {e}\nRaw: {raw[:200]}")
    
    return {}  # Unreachable but satisfies type checker


def validate_payload(data: Dict) -> None:
    """Ensure data structure is correct."""
    # Check name field exists
    if not isinstance(data.get("name"), str):
        raise ValueError("Invalid 'name' field")
    
    # Check all sections have exactly 3 items
    for section in SECTIONS:
        items = data.get(section, [])
        if (not isinstance(items, list) or 
            len(items) != 3 or 
            not all(isinstance(x, str) and x.strip() for x in items)):
            raise ValueError(f"'{section}' needs 3 non-empty strings")


def make_docx(data: Dict, output: str = "resume.docx") -> None:
    """Create single-page DOCX resume."""
    doc = Document()
    
    # Configure page margins (tight for one page)
    section = doc.sections[0]
    for margin in ["top", "bottom", "left", "right"]:
        setattr(section, f"{margin}_margin", Inches(0.5))
    
    # Add centered title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data["name"])
    run.font.name = "Poppins"
    run.font.size = Pt(22)
    
    # Create 2-column table (4 rows x 2 cols = 8 sections)
    table = doc.add_table(rows=4, cols=2)
    
    # Fill table cells with sections
    for idx, section in enumerate(SECTIONS):
        row = idx // 2  # Integer division for row
        col = idx % 2   # Modulo for column
        cell = table.cell(row, col)
        
        # Add section heading
        heading = cell.paragraphs[0]
        run = heading.add_run(section.upper().replace("_", " "))
        run.font.name = "Poppins"
        run.font.size = Pt(11)
        run.bold = True
        
        # Add bullet points
        for item in data[section]:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(0)
            
            run = p.add_run(f"• {item}")
            run.font.name = "Poppins"
            run.font.size = Pt(10)
    
    doc.save(output)


def main():
    """Main entry point."""
    while True:
        # Get user input
        job_title = input("Enter target job title: ").strip()
        if not job_title:
            continue
        industry = input("Enter industry: ").strip()
        if not industry:
            continue
        seniority = input("Enter seniority level: ").strip()
        if not seniority:
            continue
        
        # Generate content
        print("Generating content...")
        try:
            data = call_openai(job_title, industry, seniority)
            
            # Create document
            print("Creating document...")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = f"resume_{timestamp}.docx"
            make_docx(data, output_file)
            print(f"✓ Saved: {output_file}")
            
        except Exception as e:
            print(f"✗ Error: {e}")
        
        # Ask to continue
        if input("\nGenerate another? (y/n): ").lower() != 'y':
            break


if __name__ == "__main__":
    main()